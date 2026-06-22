from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = BASE_DIR / "docs"
OUT_PATH = OUT_DIR / "Matriz_de_Pruebas_SMRI.docx"

BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
DARK = "1F2937"
MUTED = "64748B"
RED = "9B1C1C"
GREEN = "166534"
GOLD = "7A5A00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_cell(cell, text, bold=False, fill=None, color=DARK, size=8.2, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph


def add_note(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [14400])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, 10, True, BLUE)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(body)
    set_run_font(run, 9.5, False, DARK)


def add_matrix_table(doc, rows, title):
    doc.add_paragraph(title, style="Heading 2")
    headers = ["ID", "Modulo", "Prior.", "Escenario y datos", "Pasos de prueba", "Resultado esperado", "Estado / evidencia"]
    widths = [820, 1450, 720, 2800, 3500, 3300, 1810]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, header_text in enumerate(headers):
        style_cell(header.cells[idx], header_text, bold=True, fill=LIGHT_BLUE, color=BLUE, size=8.3, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            style_cell(row.cells[idx], value, size=7.6, align=align)
    doc.add_paragraph()


def build_document():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SMRI | Matriz de pruebas funcionales")
    set_run_font(run, 8.5, False, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Documento de validacion del sistema de monitoreo y respuesta inteligente")
    set_run_font(run, 8, False, MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Matriz de Pruebas del Sistema SMRI")
    set_run_font(run, 22, True, BLUE)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(
        f"Validacion funcional de camaras RTSP/IP, reconocimiento facial, deteccion de eventos y evidencias | Fecha: {date.today().isoformat()}"
    )
    set_run_font(run, 10.5, False, MUTED)

    add_note(
        doc,
        "Objetivo",
        "Comprobar que el sistema detecta, registra y notifica eventos de seguridad sin comprometer la fluidez del stream. "
        "La matriz esta pensada para ejecutar pruebas controladas con una o mas camaras RTSP/IP activas.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "Criterios generales de aceptacion", 1)
    criteria = [
        "El stream debe mantenerse visible durante la prueba, sin quedarse en estado Sin senal por inferencias de IA.",
        "Cada evento valido debe crear registro, evidencia de imagen, severidad y descripcion comprensible.",
        "Los falsos positivos deben registrarse como observaciones y no considerarse aprobados.",
        "El reconocimiento facial debe conservar la identidad autorizada durante movimientos leves y no desplazar la caja del rostro.",
        "Las pruebas con objetos y animales deben realizarse en condiciones de luz similares a operacion real.",
    ]
    for item in criteria:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run, 9.5, False, DARK)

    add_heading(doc, "Ambiente y datos de prueba", 1)
    env_rows = [
        ("Camaras", "Minimo 1 camara RTSP/IP activa; recomendado probar tambien cambio entre 2 camaras."),
        ("Usuarios/rostros", "Una persona autorizada registrada; una persona no registrada; rostro con movimiento leve."),
        ("Objetos", "Celular, tijeras, cuchillo u objeto equivalente seguro para demostracion, bolso/mochila opcional."),
        ("Animales", "Gato o animal domestico controlado dentro del area visible."),
        ("Evidencias", "Verificar Registro en Vivo, listado de alertas, imagen del evento, fecha/hora, camara y descripcion."),
        ("Seguridad", "No usar objetos punzantes de forma riesgosa. Mantener pruebas de caida simuladas y controladas."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2600, 11800])
    set_repeat_table_header(table.rows[0])
    style_cell(table.cell(0, 0), "Elemento", True, LIGHT_BLUE, BLUE, 8.6, WD_ALIGN_PARAGRAPH.CENTER)
    style_cell(table.cell(0, 1), "Configuracion / dato requerido", True, LIGHT_BLUE, BLUE, 8.6)
    for label, value in env_rows:
        row = table.add_row()
        style_cell(row.cells[0], label, True, LIGHT_GRAY, DARK, 8.2)
        style_cell(row.cells[1], value, False, None, DARK, 8.2)

    test_rows_core = [
        ("RTSP-01", "Stream", "Alta", "Abrir camara RTSP/IP principal.", "Ingresar al modulo Camara y observar por 2 min.", "Video fluido; no aparece Sin senal; fecha/hora se mantiene.", "Captura del stream y log de inicio."),
        ("RTSP-02", "Stream", "Alta", "Cambiar entre dos camaras activas.", "Seleccionar camara A, luego B, luego A.", "No se congelan ambas camaras; se actualiza nombre/ubicacion.", "Captura antes/despues."),
        ("RTSP-03", "Stream", "Media", "Movimiento alto en escena.", "Moverse frente a camara por 60 s.", "Latencia aceptable; no se detiene el worker.", "Medir retraso aproximado."),
        ("FACE-01", "Rostro", "Alta", "Persona autorizada frente a camara.", "Mirar al frente, moverse leve izquierda/derecha.", "Caja sobre el rostro real; etiqueta Autorizado con nombre correcto.", "Captura con etiqueta."),
        ("FACE-02", "Rostro", "Alta", "Persona autorizada con movimiento leve.", "Girar cabeza, acercarse y alejarse.", "No cambia a No autorizado por lecturas aisladas.", "Registro sin intrusion falsa."),
        ("FACE-03", "Rostro", "Alta", "Persona no registrada.", "Ingresar al area sin registro facial.", "Tras confirmacion, alerta de persona no autorizada.", "Evento intrusion con evidencia."),
        ("FACE-04", "Rostro", "Media", "Reflejo/espejo visible.", "Moverse frente a espejo dentro del encuadre.", "No debe crear una segunda persona falsa si no hay rostro fisico en area.", "Registrar falso positivo si ocurre."),
        ("FACE-05", "Rostro", "Media", "Imagen/poster con rostro al fondo.", "Mantener imagen en pared durante 2 min.", "No debe alertar como persona no autorizada por rostro pequeno/lejos.", "Captura y log."),
    ]
    add_matrix_table(doc, test_rows_core, "Matriz A. Stream y reconocimiento facial")

    test_rows_events = [
        ("OBJ-01", "Objeto peligroso", "Alta", "Tijeras visibles cerca de persona.", "Mostrar tijeras junto a la mano por 10-15 s.", "Alerta objeto peligroso; descripcion incluye posibles cortes.", "Evento + evidencia."),
        ("OBJ-02", "Objeto peligroso", "Alta", "Cuchillo visible cerca de persona.", "Mostrar cuchillo/objeto seguro equivalente por 10-15 s.", "Alerta objeto peligroso; posible riesgo de cortes.", "Evento + evidencia."),
        ("OBJ-03", "Objeto peligroso", "Media", "Tijeras parcialmente tapadas.", "Cubrir parcialmente el objeto.", "Puede no alertar; si alerta, caja debe cubrir objeto real.", "Anotar sensibilidad."),
        ("OBJ-04", "Objeto no autorizado", "Media", "Bolso/mochila en area.", "Colocar objeto en escena.", "Alerta objeto no autorizado solo si regla esta activa.", "Evento o observacion."),
        ("PHONE-01", "Celular", "Alta", "Celular visible sin persona asociada.", "Colocar celular dentro del area visible.", "Alerta inmediata: celular no autorizado en el area.", "Evento + log PHONE."),
        ("PHONE-02", "Celular", "Alta", "Persona no identificada usando celular.", "Usar celular antes de que el rostro sea identificado.", "Alerta menor al tiempo prolongado; descripcion persona no identificada.", "Evento + evidencia."),
        ("PHONE-03", "Celular", "Alta", "Persona no autorizada usando celular.", "Persona no registrada sostiene celular.", "Alerta inmediata y evento asociado al area/persona.", "Evento + evidencia."),
        ("PHONE-04", "Celular", "Media", "Persona autorizada usa celular.", "Persona autorizada sostiene celular por menos de 10 s.", "Log PHONE aparece; no debe crear alerta prolongada antes del umbral.", "Log en vivo."),
        ("PHONE-05", "Celular", "Alta", "Persona autorizada usa celular prolongado.", "Mantener celular por mas de 10 s.", "Alerta de uso prolongado con identidad si esta disponible.", "Evento + evidencia."),
        ("ANIMAL-01", "Animal", "Alta", "Gato en area monitoreada.", "Ingresar gato controlado al encuadre.", "Alerta acceso no autorizado: gato o animal posible perro/gato.", "Evento + evidencia."),
        ("ANIMAL-02", "Animal", "Media", "Animal parcialmente visible.", "Mostrar animal solo parcialmente.", "Sistema debe alertar si confianza supera umbral; registrar si lo confunde.", "Observacion de clase/confianza."),
    ]
    add_matrix_table(doc, test_rows_events, "Matriz B. Eventos visuales: objetos, celular y animales")

    test_rows_pose = [
        ("POSE-01", "Caida", "Alta", "Persona de pie con movimiento normal.", "Caminar, girar, sentarse y levantarse.", "No debe generar posible caida.", "Log sin alerta de caida."),
        ("POSE-02", "Caida", "Alta", "Persona inclinada brevemente.", "Agacharse a recoger un objeto por menos de 3 s.", "No debe alertar si no permanece en postura de caida.", "Sin evento."),
        ("POSE-03", "Caida", "Alta", "Caida simulada controlada.", "Acostarse o quedar en piso/postura horizontal por mas de 3 s.", "Alerta posible caida con persona identificada si aplica.", "Evento + evidencia."),
        ("POSE-04", "Caida", "Media", "Persona reflejada en espejo.", "Moverse frente a espejo.", "No debe duplicar persona ni disparar caida por reflejo.", "Captura y observacion."),
        ("POSE-05", "Multipersona", "Alta", "Dos personas visibles.", "Ingresar dos personas, una autorizada y otra no.", "Sistema debe mantener detecciones y asociar eventos al sujeto correcto cuando sea posible.", "Captura con cajas."),
        ("POSE-06", "Multipersona", "Media", "Persona + animal + objeto.", "Persona autorizada, gato y tijeras en escena.", "Debe detectar combinacion: persona, animal no autorizado y objeto peligroso si visibles.", "Eventos esperados y evidencia."),
    ]
    add_matrix_table(doc, test_rows_pose, "Matriz C. Pose, caidas y escenarios combinados")

    test_rows_evidence = [
        ("EV-01", "Evidencia", "Alta", "Alerta valida generada.", "Abrir detalle de evento.", "Imagen corresponde al momento del evento; caja no esta desplazada.", "Comparar evidencia vs escena."),
        ("EV-02", "Evidencia", "Alta", "Descripcion del evento.", "Revisar tarjeta/listado.", "Incluye tipo, categoria, severidad, confianza y camara.", "Captura del detalle."),
        ("EV-03", "Logs", "Media", "Registro en vivo.", "Provocar celular, animal y objeto peligroso.", "Log muestra PHONE/ALERTA/OBJ de forma comprensible.", "Captura del log."),
        ("EV-04", "Correo", "Media", "Evento que envia notificacion.", "Generar alerta critica.", "Estado de correo SENT, FAILED o SKIPPED segun configuracion.", "Estado en detalle."),
        ("EV-05", "Reportes", "Media", "Informe creado.", "Consultar modulo de informes.", "Existe informe asociado al SecurityEvent.", "ID de evento/informe."),
        ("PERF-01", "Rendimiento", "Alta", "Operacion 10 minutos.", "Dejar camara activa con movimiento normal.", "Sin detencion del stream; latencia estable.", "Tiempo observado."),
        ("PERF-02", "Rendimiento", "Alta", "Dos camaras activas/cambio frecuente.", "Cambiar entre camaras durante 5 min.", "No se detienen ambas camaras; workers se recuperan.", "Capturas y logs."),
        ("PERF-03", "Rendimiento", "Media", "Escena con mucho movimiento.", "Mover personas/objetos durante 2 min.", "Latencia no crece progresivamente.", "Medicion inicial/final."),
    ]
    add_matrix_table(doc, test_rows_evidence, "Matriz D. Evidencias, logs y rendimiento")

    add_heading(doc, "Plantilla de registro de ejecucion", 1)
    add_note(
        doc,
        "Uso recomendado",
        "Completar una fila por cada caso ejecutado. Si el resultado falla, adjuntar captura, hora exacta, camara, confianza mostrada y descripcion observada.",
        fill="FFF7E6",
    )
    headers = ["ID prueba", "Fecha/hora", "Camara", "Ejecutor", "Resultado", "Observaciones / evidencia"]
    widths = [1300, 1700, 1700, 1700, 1300, 6700]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header_text in enumerate(headers):
        style_cell(table.rows[0].cells[idx], header_text, True, LIGHT_BLUE, BLUE, 8.5, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(12):
        row = table.add_row()
        for idx in range(len(headers)):
            style_cell(row.cells[idx], "", False, None, DARK, 8.2)

    add_heading(doc, "Estados sugeridos", 1)
    states = [
        ("Aprobado", "El resultado observado coincide con el esperado y la evidencia es correcta."),
        ("Aprobado con observacion", "La funcion responde, pero hay detalle menor de precision, texto, tiempo o evidencia."),
        ("Fallido", "No detecta, detecta incorrectamente, crea falso positivo critico o afecta la fluidez."),
        ("No ejecutado", "No fue posible ejecutar por ambiente, datos, camara o restriccion de seguridad."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2600, 11800])
    set_repeat_table_header(table.rows[0])
    style_cell(table.cell(0, 0), "Estado", True, LIGHT_BLUE, BLUE, 8.5, WD_ALIGN_PARAGRAPH.CENTER)
    style_cell(table.cell(0, 1), "Criterio", True, LIGHT_BLUE, BLUE, 8.5)
    for status, detail in states:
        row = table.add_row()
        fill = None
        color = DARK
        if status == "Aprobado":
            color = GREEN
        elif status == "Fallido":
            color = RED
        elif status == "Aprobado con observacion":
            color = GOLD
        style_cell(row.cells[0], status, True, fill, color, 8.2)
        style_cell(row.cells[1], detail, False, fill, DARK, 8.2)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
